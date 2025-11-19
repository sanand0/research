"""
Unit tests for File Renamer module.
"""
import os
import tempfile
import shutil
import unittest
from pathlib import Path

from file_renamer import FileRenamer, SuffixType, PREDEFINED_PREFIXES


class TestFileRenamer(unittest.TestCase):
    """Test cases for FileRenamer class."""

    def setUp(self):
        """Set up test fixtures."""
        self.test_dir = tempfile.mkdtemp()
        self.renamer = FileRenamer()

    def tearDown(self):
        """Clean up test fixtures."""
        if os.path.exists(self.test_dir):
            shutil.rmtree(self.test_dir)

    def create_test_files(self, count):
        """
        Helper method to create test files.

        Args:
            count: Number of files to create

        Returns:
            List of file paths
        """
        files = []
        for i in range(count):
            filepath = os.path.join(self.test_dir, f"file_{i}.txt")
            with open(filepath, 'w') as f:
                f.write(f"Content of file {i}")
            files.append(filepath)
        return files

    def test_pause_resume(self):
        """Test pause and resume functionality."""
        self.assertFalse(self.renamer.is_paused)

        self.renamer.pause()
        self.assertTrue(self.renamer.is_paused)

        self.renamer.resume()
        self.assertFalse(self.renamer.is_paused)

    def test_stop(self):
        """Test stop functionality."""
        self.assertFalse(self.renamer.is_stopped)

        self.renamer.stop()
        self.assertTrue(self.renamer.is_stopped)
        self.assertFalse(self.renamer.is_paused)

    def test_reset(self):
        """Test reset functionality."""
        self.renamer.pause()
        self.renamer.stop()

        self.renamer.reset()
        self.assertFalse(self.renamer.is_paused)
        self.assertFalse(self.renamer.is_stopped)

    def test_rename_files_numeric(self):
        """Test renaming with numeric suffix."""
        files = self.create_test_files(3)
        output_dir = os.path.join(self.test_dir, "output")

        results = self.renamer.rename_files(
            files,
            output_dir,
            "Chapter",
            SuffixType.NUMERIC
        )

        self.assertEqual(len(results), 3)
        self.assertTrue(os.path.exists(os.path.join(output_dir, "Chapter 1.txt")))
        self.assertTrue(os.path.exists(os.path.join(output_dir, "Chapter 2.txt")))
        self.assertTrue(os.path.exists(os.path.join(output_dir, "Chapter 3.txt")))

    def test_rename_files_alpha_lower(self):
        """Test renaming with lowercase alphabetic suffix."""
        files = self.create_test_files(3)
        output_dir = os.path.join(self.test_dir, "output")

        results = self.renamer.rename_files(
            files,
            output_dir,
            "Section",
            SuffixType.ALPHA_LOWER
        )

        self.assertEqual(len(results), 3)
        self.assertTrue(os.path.exists(os.path.join(output_dir, "Section a.txt")))
        self.assertTrue(os.path.exists(os.path.join(output_dir, "Section b.txt")))
        self.assertTrue(os.path.exists(os.path.join(output_dir, "Section c.txt")))

    def test_rename_files_alpha_upper(self):
        """Test renaming with uppercase alphabetic suffix."""
        files = self.create_test_files(3)
        output_dir = os.path.join(self.test_dir, "output")

        results = self.renamer.rename_files(
            files,
            output_dir,
            "Part",
            SuffixType.ALPHA_UPPER
        )

        self.assertEqual(len(results), 3)
        self.assertTrue(os.path.exists(os.path.join(output_dir, "Part A.txt")))
        self.assertTrue(os.path.exists(os.path.join(output_dir, "Part B.txt")))
        self.assertTrue(os.path.exists(os.path.join(output_dir, "Part C.txt")))

    def test_rename_files_roman_lower(self):
        """Test renaming with lowercase Roman numeral suffix."""
        files = self.create_test_files(3)
        output_dir = os.path.join(self.test_dir, "output")

        results = self.renamer.rename_files(
            files,
            output_dir,
            "Chapter",
            SuffixType.ROMAN_LOWER
        )

        self.assertEqual(len(results), 3)
        self.assertTrue(os.path.exists(os.path.join(output_dir, "Chapter i.txt")))
        self.assertTrue(os.path.exists(os.path.join(output_dir, "Chapter ii.txt")))
        self.assertTrue(os.path.exists(os.path.join(output_dir, "Chapter iii.txt")))

    def test_rename_files_roman_upper(self):
        """Test renaming with uppercase Roman numeral suffix."""
        files = self.create_test_files(3)
        output_dir = os.path.join(self.test_dir, "output")

        results = self.renamer.rename_files(
            files,
            output_dir,
            "Chapter",
            SuffixType.ROMAN_UPPER
        )

        self.assertEqual(len(results), 3)
        self.assertTrue(os.path.exists(os.path.join(output_dir, "Chapter I.txt")))
        self.assertTrue(os.path.exists(os.path.join(output_dir, "Chapter II.txt")))
        self.assertTrue(os.path.exists(os.path.join(output_dir, "Chapter III.txt")))

    def test_rename_with_custom_start_number(self):
        """Test renaming with custom start number."""
        files = self.create_test_files(3)
        output_dir = os.path.join(self.test_dir, "output")

        results = self.renamer.rename_files(
            files,
            output_dir,
            "Chapter",
            SuffixType.NUMERIC,
            start_number=10
        )

        self.assertEqual(len(results), 3)
        self.assertTrue(os.path.exists(os.path.join(output_dir, "Chapter 10.txt")))
        self.assertTrue(os.path.exists(os.path.join(output_dir, "Chapter 11.txt")))
        self.assertTrue(os.path.exists(os.path.join(output_dir, "Chapter 12.txt")))

    def test_rename_empty_file_list(self):
        """Test renaming with empty file list."""
        output_dir = os.path.join(self.test_dir, "output")

        with self.assertRaises(ValueError):
            self.renamer.rename_files(
                [],
                output_dir,
                "Chapter",
                SuffixType.NUMERIC
            )

    def test_rename_empty_prefix(self):
        """Test renaming with empty prefix."""
        files = self.create_test_files(1)
        output_dir = os.path.join(self.test_dir, "output")

        with self.assertRaises(ValueError):
            self.renamer.rename_files(
                files,
                output_dir,
                "",
                SuffixType.NUMERIC
            )

    def test_rename_nonexistent_file(self):
        """Test renaming with nonexistent file."""
        output_dir = os.path.join(self.test_dir, "output")

        with self.assertRaises(FileNotFoundError):
            self.renamer.rename_files(
                ["nonexistent.txt"],
                output_dir,
                "Chapter",
                SuffixType.NUMERIC
            )

    def test_rename_with_different_extensions(self):
        """Test renaming files with different extensions."""
        # Create files with different extensions
        files = []
        for ext in ['.txt', '.docx', '.pdf']:
            filepath = os.path.join(self.test_dir, f"file{ext}")
            with open(filepath, 'w') as f:
                f.write("content")
            files.append(filepath)

        output_dir = os.path.join(self.test_dir, "output")

        results = self.renamer.rename_files(
            files,
            output_dir,
            "Document",
            SuffixType.NUMERIC
        )

        self.assertEqual(len(results), 3)
        self.assertTrue(os.path.exists(os.path.join(output_dir, "Document 1.txt")))
        self.assertTrue(os.path.exists(os.path.join(output_dir, "Document 2.docx")))
        self.assertTrue(os.path.exists(os.path.join(output_dir, "Document 3.pdf")))

    def test_rename_preserves_file_content(self):
        """Test that renaming preserves file content."""
        test_content = "Important content that should be preserved"
        filepath = os.path.join(self.test_dir, "original.txt")
        with open(filepath, 'w') as f:
            f.write(test_content)

        output_dir = os.path.join(self.test_dir, "output")

        results = self.renamer.rename_files(
            [filepath],
            output_dir,
            "Renamed",
            SuffixType.NUMERIC
        )

        new_file = os.path.join(output_dir, "Renamed 1.txt")
        with open(new_file, 'r') as f:
            content = f.read()

        self.assertEqual(content, test_content)

    def test_output_directory_creation(self):
        """Test that output directory is created if it doesn't exist."""
        files = self.create_test_files(1)
        output_dir = os.path.join(self.test_dir, "new_output_dir")

        self.assertFalse(os.path.exists(output_dir))

        self.renamer.rename_files(
            files,
            output_dir,
            "File",
            SuffixType.NUMERIC
        )

        self.assertTrue(os.path.exists(output_dir))

    def test_rename_with_progress_callback(self):
        """Test renaming with progress callback."""
        files = self.create_test_files(3)
        output_dir = os.path.join(self.test_dir, "output")

        progress_updates = []

        def progress_callback(percent, message):
            progress_updates.append((percent, message))

        self.renamer.rename_files(
            files,
            output_dir,
            "Chapter",
            SuffixType.NUMERIC,
            progress_callback=progress_callback
        )

        # Verify progress updates were made
        self.assertGreater(len(progress_updates), 0)
        self.assertEqual(progress_updates[-1][0], 100)

    def test_number_to_alpha_basic(self):
        """Test number to alpha conversion."""
        self.assertEqual(self.renamer._number_to_alpha(1), "a")
        self.assertEqual(self.renamer._number_to_alpha(26), "z")
        self.assertEqual(self.renamer._number_to_alpha(27), "aa")
        self.assertEqual(self.renamer._number_to_alpha(52), "az")

    def test_number_to_alpha_uppercase(self):
        """Test number to alpha conversion with uppercase."""
        self.assertEqual(self.renamer._number_to_alpha(1, uppercase=True), "A")
        self.assertEqual(self.renamer._number_to_alpha(26, uppercase=True), "Z")
        self.assertEqual(self.renamer._number_to_alpha(27, uppercase=True), "AA")

    def test_number_to_roman_basic(self):
        """Test number to Roman numeral conversion."""
        self.assertEqual(self.renamer._number_to_roman(1), "I")
        self.assertEqual(self.renamer._number_to_roman(2), "II")
        self.assertEqual(self.renamer._number_to_roman(3), "III")
        self.assertEqual(self.renamer._number_to_roman(4), "IV")
        self.assertEqual(self.renamer._number_to_roman(5), "V")
        self.assertEqual(self.renamer._number_to_roman(9), "IX")
        self.assertEqual(self.renamer._number_to_roman(10), "X")

    def test_number_to_roman_complex(self):
        """Test complex Roman numeral conversions."""
        self.assertEqual(self.renamer._number_to_roman(49), "XLIX")
        self.assertEqual(self.renamer._number_to_roman(99), "XCIX")
        self.assertEqual(self.renamer._number_to_roman(444), "CDXLIV")
        self.assertEqual(self.renamer._number_to_roman(1994), "MCMXCIV")

    def test_number_to_roman_invalid_range(self):
        """Test Roman numeral conversion with invalid range."""
        with self.assertRaises(ValueError):
            self.renamer._number_to_roman(0)

        with self.assertRaises(ValueError):
            self.renamer._number_to_roman(4000)

    def test_generate_suffix_all_types(self):
        """Test suffix generation for all types."""
        test_number = 5

        self.assertEqual(
            self.renamer._generate_suffix(test_number, SuffixType.NUMERIC),
            "5"
        )
        self.assertEqual(
            self.renamer._generate_suffix(test_number, SuffixType.ALPHA_LOWER),
            "e"
        )
        self.assertEqual(
            self.renamer._generate_suffix(test_number, SuffixType.ALPHA_UPPER),
            "E"
        )
        self.assertEqual(
            self.renamer._generate_suffix(test_number, SuffixType.ROMAN_LOWER),
            "v"
        )
        self.assertEqual(
            self.renamer._generate_suffix(test_number, SuffixType.ROMAN_UPPER),
            "V"
        )

    def test_predefined_prefixes(self):
        """Test that predefined prefixes are available."""
        self.assertIn("Chapter", PREDEFINED_PREFIXES)
        self.assertIn("Section", PREDEFINED_PREFIXES)
        self.assertIn("Part", PREDEFINED_PREFIXES)
        self.assertIn("Figure", PREDEFINED_PREFIXES)
        self.assertIn("Table", PREDEFINED_PREFIXES)
        self.assertIn("Equation", PREDEFINED_PREFIXES)


class TestFileRenamerIntegration(unittest.TestCase):
    """Integration tests for FileRenamer."""

    def setUp(self):
        """Set up test fixtures."""
        self.test_dir = tempfile.mkdtemp()
        self.renamer = FileRenamer()

    def tearDown(self):
        """Clean up test fixtures."""
        if os.path.exists(self.test_dir):
            shutil.rmtree(self.test_dir)

    def test_full_workflow_with_word_files(self):
        """Test complete workflow with Word files."""
        from docx import Document

        # Create test Word files
        files = []
        for i in range(5):
            doc = Document()
            doc.add_paragraph(f"Chapter {i+1} content")
            filepath = os.path.join(self.test_dir, f"chapter_{i+1}.docx")
            doc.save(filepath)
            files.append(filepath)

        output_dir = os.path.join(self.test_dir, "renamed")

        # Rename the files
        results = self.renamer.rename_files(
            files,
            output_dir,
            "Chapter",
            SuffixType.ROMAN_UPPER,
            start_number=1
        )

        # Verify results
        self.assertEqual(len(results), 5)
        self.assertTrue(os.path.exists(os.path.join(output_dir, "Chapter I.docx")))
        self.assertTrue(os.path.exists(os.path.join(output_dir, "Chapter V.docx")))

        # Verify files are still valid Word documents
        for _, new_path in results:
            doc = Document(new_path)
            self.assertGreater(len(doc.paragraphs), 0)

    def test_large_batch_renaming(self):
        """Test renaming a large batch of files."""
        # Create 100 test files
        files = []
        for i in range(100):
            filepath = os.path.join(self.test_dir, f"file_{i:03d}.txt")
            with open(filepath, 'w') as f:
                f.write(f"Content {i}")
            files.append(filepath)

        output_dir = os.path.join(self.test_dir, "output")

        # Rename all files
        results = self.renamer.rename_files(
            files,
            output_dir,
            "Item",
            SuffixType.NUMERIC
        )

        # Verify all files were renamed
        self.assertEqual(len(results), 100)
        self.assertTrue(os.path.exists(os.path.join(output_dir, "Item 1.txt")))
        self.assertTrue(os.path.exists(os.path.join(output_dir, "Item 100.txt")))


if __name__ == '__main__':
    unittest.main()
