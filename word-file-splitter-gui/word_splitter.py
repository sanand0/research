"""
Word File Splitter Module
Handles splitting Word documents based on delimiters.
"""
import os
import shutil
from pathlib import Path
from typing import List, Tuple, Callable, Optional
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph


class WordSplitter:
    """Handles splitting Word documents based on delimiters."""

    def __init__(self, delimiter: str = "***"):
        """
        Initialize the WordSplitter.

        Args:
            delimiter: The text delimiter to split on (default: ***)
        """
        self.delimiter = delimiter
        self.is_paused = False
        self.is_stopped = False

    def pause(self):
        """Pause the splitting process."""
        self.is_paused = True

    def resume(self):
        """Resume the splitting process."""
        self.is_paused = False

    def stop(self):
        """Stop the splitting process."""
        self.is_stopped = True
        self.is_paused = False

    def reset(self):
        """Reset the state for a new operation."""
        self.is_paused = False
        self.is_stopped = False

    def split_single_file(
        self,
        input_path: str,
        output_dir: str,
        progress_callback: Optional[Callable[[int, str], None]] = None
    ) -> List[str]:
        """
        Split a single Word file based on the delimiter.

        Args:
            input_path: Path to the input Word file
            output_dir: Directory to save the split files
            progress_callback: Optional callback function(progress_percent, message)

        Returns:
            List of created file paths

        Raises:
            FileNotFoundError: If input file doesn't exist
            ValueError: If file is not a valid Word document
        """
        self.reset()

        if not os.path.exists(input_path):
            raise FileNotFoundError(f"Input file not found: {input_path}")

        if not input_path.lower().endswith(('.docx', '.doc')):
            raise ValueError("Input file must be a Word document (.docx or .doc)")

        # Create output directory if it doesn't exist
        os.makedirs(output_dir, exist_ok=True)

        # Create temporary copy to preserve notes
        temp_path = os.path.join(output_dir, f"_temp_{os.path.basename(input_path)}")
        shutil.copy2(input_path, temp_path)

        try:
            # Load the document
            doc = Document(temp_path)

            if progress_callback:
                progress_callback(10, "Document loaded, finding split points...")

            # Find all delimiter positions
            split_indices = self._find_split_indices(doc)

            if not split_indices:
                if progress_callback:
                    progress_callback(100, "No delimiters found. File not split.")
                return []

            if progress_callback:
                progress_callback(20, f"Found {len(split_indices)} split points...")

            # Split the document
            output_files = self._split_document(
                doc,
                split_indices,
                input_path,
                output_dir,
                progress_callback
            )

            if progress_callback and not self.is_stopped:
                progress_callback(100, f"Successfully created {len(output_files)} files")

            return output_files

        finally:
            # Clean up temporary file
            if os.path.exists(temp_path):
                os.remove(temp_path)

    def split_batch(
        self,
        input_dir: str,
        output_dir: str,
        progress_callback: Optional[Callable[[int, str], None]] = None
    ) -> dict:
        """
        Split multiple Word files in a directory.

        Args:
            input_dir: Directory containing Word files to split
            output_dir: Directory to save the split files
            progress_callback: Optional callback function(progress_percent, message)

        Returns:
            Dictionary with results: {'success': [...], 'failed': [...]}
        """
        self.reset()

        if not os.path.exists(input_dir):
            raise FileNotFoundError(f"Input directory not found: {input_dir}")

        # Find all Word files
        word_files = []
        for ext in ['*.docx', '*.doc']:
            word_files.extend(Path(input_dir).glob(ext))

        if not word_files:
            if progress_callback:
                progress_callback(100, "No Word files found in input directory")
            return {'success': [], 'failed': []}

        results = {'success': [], 'failed': []}
        total_files = len(word_files)

        for idx, file_path in enumerate(word_files):
            if self.is_stopped:
                if progress_callback:
                    progress_callback(100, "Processing stopped by user")
                break

            while self.is_paused:
                if self.is_stopped:
                    break
                continue

            try:
                file_name = file_path.name
                if progress_callback:
                    progress_callback(
                        int((idx / total_files) * 100),
                        f"Processing {file_name}..."
                    )

                # Create subdirectory for this file's output
                file_output_dir = os.path.join(
                    output_dir,
                    file_path.stem
                )

                output_files = self.split_single_file(
                    str(file_path),
                    file_output_dir
                )

                results['success'].append({
                    'input': str(file_path),
                    'outputs': output_files
                })

            except Exception as e:
                results['failed'].append({
                    'input': str(file_path),
                    'error': str(e)
                })

        if progress_callback and not self.is_stopped:
            progress_callback(
                100,
                f"Batch complete: {len(results['success'])} succeeded, "
                f"{len(results['failed'])} failed"
            )

        return results

    def _find_split_indices(self, doc: Document) -> List[int]:
        """
        Find all paragraph indices where the delimiter appears.

        Args:
            doc: The Document object

        Returns:
            List of paragraph indices containing the delimiter
        """
        split_indices = []

        for idx, paragraph in enumerate(doc.paragraphs):
            if self.delimiter in paragraph.text:
                split_indices.append(idx)

        return split_indices

    def _split_document(
        self,
        doc: Document,
        split_indices: List[int],
        original_path: str,
        output_dir: str,
        progress_callback: Optional[Callable[[int, str], None]] = None
    ) -> List[str]:
        """
        Split the document at the specified indices.

        Args:
            doc: The Document object
            split_indices: List of paragraph indices to split at
            original_path: Original file path (for naming)
            output_dir: Output directory
            progress_callback: Optional progress callback

        Returns:
            List of created file paths
        """
        output_files = []
        base_name = Path(original_path).stem

        # Create sections: [0, split1), [split1, split2), ..., [splitN, end)
        sections = []
        prev_idx = 0

        for split_idx in split_indices:
            sections.append((prev_idx, split_idx))
            prev_idx = split_idx

        # Add final section
        sections.append((prev_idx, len(doc.paragraphs)))

        total_sections = len(sections)
        base_progress = 20  # We're at 20% when we start splitting
        progress_range = 80  # We have 80% progress left

        for section_num, (start_idx, end_idx) in enumerate(sections):
            if self.is_stopped:
                break

            while self.is_paused:
                if self.is_stopped:
                    break
                continue

            # Skip empty sections
            if start_idx >= end_idx:
                continue

            # Create new document for this section
            new_doc = Document()

            # Copy the section's paragraphs
            for idx in range(start_idx, end_idx):
                if idx < len(doc.paragraphs):
                    para = doc.paragraphs[idx]

                    # Skip the delimiter paragraph itself
                    if self.delimiter in para.text and idx == start_idx:
                        continue

                    # Copy paragraph with formatting
                    self._copy_paragraph(para, new_doc)

            # Save the section
            output_filename = f"{base_name}_part_{section_num + 1:03d}.docx"
            output_path = os.path.join(output_dir, output_filename)
            new_doc.save(output_path)
            output_files.append(output_path)

            if progress_callback:
                progress = base_progress + int(
                    ((section_num + 1) / total_sections) * progress_range
                )
                progress_callback(
                    progress,
                    f"Created {output_filename}"
                )

        return output_files

    def _copy_paragraph(self, source_para: Paragraph, target_doc: Document):
        """
        Copy a paragraph from source to target document with formatting.

        Args:
            source_para: Source paragraph
            target_doc: Target document
        """
        new_para = target_doc.add_paragraph(source_para.text, source_para.style)

        # Copy paragraph formatting
        new_para.paragraph_format.alignment = source_para.paragraph_format.alignment
        new_para.paragraph_format.left_indent = source_para.paragraph_format.left_indent
        new_para.paragraph_format.right_indent = source_para.paragraph_format.right_indent
        new_para.paragraph_format.first_line_indent = source_para.paragraph_format.first_line_indent
        new_para.paragraph_format.space_before = source_para.paragraph_format.space_before
        new_para.paragraph_format.space_after = source_para.paragraph_format.space_after
        new_para.paragraph_format.line_spacing = source_para.paragraph_format.line_spacing

        # Copy run formatting
        if source_para.runs and new_para.runs:
            for src_run, tgt_run in zip(source_para.runs, new_para.runs):
                tgt_run.bold = src_run.bold
                tgt_run.italic = src_run.italic
                tgt_run.underline = src_run.underline
                tgt_run.font.name = src_run.font.name
                tgt_run.font.size = src_run.font.size
