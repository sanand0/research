"""
Script to create sample Word file fixtures for testing.
"""
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def create_sample_book():
    """Create a sample book with chapters for splitting."""
    doc = Document()

    # Title
    title = doc.add_paragraph("Sample Book for Testing", style='Title')
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Front Matter
    doc.add_paragraph("Front Matter", style='Heading 1')
    doc.add_paragraph(
        "This is the front matter of the book. It contains information "
        "about the book, acknowledgments, and other preliminary content."
    )
    doc.add_paragraph(
        "Lorem ipsum dolor sit amet, consectetur adipiscing elit. "
        "Sed do eiusmod tempor incididunt ut labore et dolore magna aliqua."
    )

    # Delimiter
    doc.add_paragraph("***")

    # Chapter 1
    doc.add_paragraph("Chapter 1: Introduction", style='Heading 1')
    doc.add_paragraph(
        "This is the first chapter of the book. It introduces the main "
        "concepts and sets the stage for what's to come."
    )
    doc.add_paragraph(
        "In this chapter, we will cover the following topics:"
    )
    doc.add_paragraph("• Background and motivation", style='List Bullet')
    doc.add_paragraph("• Research objectives", style='List Bullet')
    doc.add_paragraph("• Thesis structure", style='List Bullet')
    doc.add_paragraph(
        "The introduction provides a comprehensive overview of the research "
        "area and establishes the context for the subsequent chapters."
    )

    # Delimiter
    doc.add_paragraph("***")

    # Chapter 2
    doc.add_paragraph("Chapter 2: Literature Review", style='Heading 1')
    doc.add_paragraph(
        "This chapter reviews the existing literature related to our research topic."
    )
    doc.add_paragraph("Section 2.1: Historical Context", style='Heading 2')
    doc.add_paragraph(
        "The field has a rich history dating back several decades. "
        "Early pioneers established foundational principles that continue "
        "to influence modern approaches."
    )
    doc.add_paragraph("Section 2.2: Current State of the Art", style='Heading 2')
    doc.add_paragraph(
        "Recent developments have significantly advanced our understanding. "
        "Modern techniques leverage computational power and sophisticated "
        "algorithms to achieve unprecedented results."
    )

    # Delimiter
    doc.add_paragraph("***")

    # Chapter 3
    doc.add_paragraph("Chapter 3: Methodology", style='Heading 1')
    doc.add_paragraph(
        "This chapter describes the methods and approaches used in this research."
    )
    doc.add_paragraph(
        "We employed a mixed-methods approach combining quantitative and "
        "qualitative techniques. The research design was carefully constructed "
        "to ensure validity and reliability of results."
    )
    doc.add_paragraph("Section 3.1: Data Collection", style='Heading 2')
    doc.add_paragraph(
        "Data was collected from multiple sources using standardized protocols. "
        "All procedures were approved by the institutional review board."
    )

    # Delimiter
    doc.add_paragraph("***")

    # Chapter 4
    doc.add_paragraph("Chapter 4: Results", style='Heading 1')
    doc.add_paragraph(
        "This chapter presents the findings from our research."
    )
    doc.add_paragraph(
        "The results demonstrate significant relationships between key variables. "
        "Statistical analysis revealed patterns that support our initial hypotheses."
    )
    doc.add_paragraph(
        "Key findings include:"
    )
    doc.add_paragraph("1. Finding one with detailed explanation", style='List Number')
    doc.add_paragraph("2. Finding two with supporting evidence", style='List Number')
    doc.add_paragraph("3. Finding three with implications", style='List Number')

    # Delimiter
    doc.add_paragraph("***")

    # Chapter 5
    doc.add_paragraph("Chapter 5: Discussion and Conclusion", style='Heading 1')
    doc.add_paragraph(
        "This final chapter discusses the implications of our findings and "
        "provides concluding remarks."
    )
    doc.add_paragraph(
        "Our research contributes to the field by demonstrating new approaches "
        "and validating theoretical frameworks. The results have important "
        "implications for both theory and practice."
    )
    doc.add_paragraph("Section 5.1: Limitations", style='Heading 2')
    doc.add_paragraph(
        "As with any research, this study has certain limitations that should "
        "be acknowledged. Future work could address these limitations."
    )
    doc.add_paragraph("Section 5.2: Future Directions", style='Heading 2')
    doc.add_paragraph(
        "Several promising avenues for future research have emerged from this work. "
        "We encourage other researchers to build upon these findings."
    )

    return doc


def create_short_sample():
    """Create a shorter sample document."""
    doc = Document()

    doc.add_paragraph("Short Document Example", style='Title')

    doc.add_paragraph("Part 1: Beginning", style='Heading 1')
    doc.add_paragraph("This is the beginning of the document.")

    doc.add_paragraph("***")

    doc.add_paragraph("Part 2: Middle", style='Heading 1')
    doc.add_paragraph("This is the middle section.")

    doc.add_paragraph("***")

    doc.add_paragraph("Part 3: End", style='Heading 1')
    doc.add_paragraph("This is the end of the document.")

    return doc


def create_custom_delimiter_sample():
    """Create a sample with custom delimiter."""
    doc = Document()

    doc.add_paragraph("Document with Custom Delimiter", style='Title')

    doc.add_paragraph("Section A", style='Heading 1')
    doc.add_paragraph("First section content.")

    doc.add_paragraph("###")

    doc.add_paragraph("Section B", style='Heading 1')
    doc.add_paragraph("Second section content.")

    doc.add_paragraph("###")

    doc.add_paragraph("Section C", style='Heading 1')
    doc.add_paragraph("Third section content.")

    return doc


def create_files_for_renaming():
    """Create sample files for renaming tests."""
    files_info = [
        ("chapter_draft_01.txt", "Draft content for chapter 1"),
        ("chapter_draft_02.txt", "Draft content for chapter 2"),
        ("chapter_draft_03.txt", "Draft content for chapter 3"),
        ("figure_img_a.txt", "Placeholder for figure file"),
        ("table_data_1.txt", "Placeholder for table file"),
    ]
    return files_info


def main():
    """Create all sample fixtures."""
    # Create fixtures directory
    fixtures_dir = "fixtures"
    os.makedirs(fixtures_dir, exist_ok=True)

    # Create sample book
    print("Creating sample book...")
    book_doc = create_sample_book()
    book_path = os.path.join(fixtures_dir, "sample_book.docx")
    book_doc.save(book_path)
    print(f"Created: {book_path}")

    # Create short sample
    print("Creating short sample...")
    short_doc = create_short_sample()
    short_path = os.path.join(fixtures_dir, "short_sample.docx")
    short_doc.save(short_path)
    print(f"Created: {short_path}")

    # Create custom delimiter sample
    print("Creating custom delimiter sample...")
    custom_doc = create_custom_delimiter_sample()
    custom_path = os.path.join(fixtures_dir, "custom_delimiter_sample.docx")
    custom_doc.save(custom_path)
    print(f"Created: {custom_path}")

    # Create files for renaming
    print("Creating sample files for renaming...")
    rename_dir = os.path.join(fixtures_dir, "files_to_rename")
    os.makedirs(rename_dir, exist_ok=True)

    for filename, content in create_files_for_renaming():
        filepath = os.path.join(rename_dir, filename)
        with open(filepath, 'w') as f:
            f.write(content)
        print(f"Created: {filepath}")

    # Create batch processing samples
    print("Creating batch processing samples...")
    batch_dir = os.path.join(fixtures_dir, "batch_samples")
    os.makedirs(batch_dir, exist_ok=True)

    for i in range(3):
        doc = Document()
        doc.add_paragraph(f"Document {i+1}", style='Title')
        doc.add_paragraph("Part 1", style='Heading 1')
        doc.add_paragraph(f"Content of part 1 in document {i+1}")
        doc.add_paragraph("***")
        doc.add_paragraph("Part 2", style='Heading 1')
        doc.add_paragraph(f"Content of part 2 in document {i+1}")

        filepath = os.path.join(batch_dir, f"document_{i+1}.docx")
        doc.save(filepath)
        print(f"Created: {filepath}")

    print("\nAll fixtures created successfully!")
    print(f"Fixtures location: {os.path.abspath(fixtures_dir)}")


if __name__ == "__main__":
    main()
