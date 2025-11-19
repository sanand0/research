"""
Unit tests for Word Splitter module.
"""
import os
import tempfile
import shutil
import unittest
from pathlib import Path
from docx import Document

from word_splitter import WordSplitter


class TestWordSplitter(unittest.TestCase):
    """Test cases for WordSplitter class."""

    def setUp(self):
        """Set up test fixtures."""
        self.test_dir = tempfile.mkdtemp()
        self.splitter = WordSplitter()

    def tearDown(self):
        """Clean up test fixtures."""
        if os.path.exists(self.test_dir):
            shutil.rmtree(self.test_dir)

    def create_test_document(self, filename, paragraphs):
        """
        Helper method to create a test Word document.

        Args:
            filename: Name of the file
            paragraphs: List of paragraph texts

        Returns:
            Path to the created document
        """
        doc = Document()
        for text in paragraphs:
            doc.add_paragraph(text)

        filepath = os.path.join(self.test_dir, filename)
        doc.save(filepath)
        return filepath

    def test_init_default_delimiter(self):
        """Test initialization with default delimiter."""
        splitter = WordSplitter()
        self.assertEqual(splitter.delimiter, "***")

    def test_init_custom_delimiter(self):
        """Test initialization with custom delimiter."""
        splitter = WordSplitter(delimiter="###")
        self.assertEqual(splitter.delimiter, "###")

    def test_pause_resume(self):
        """Test pause and resume functionality."""
        self.assertFalse(self.splitter.is_paused)

        self.splitter.pause()
        self.assertTrue(self.splitter.is_paused)

        self.splitter.resume()
        self.assertFalse(self.splitter.is_paused)

    def test_stop(self):
        """Test stop functionality."""
        self.assertFalse(self.splitter.is_stopped)

        self.splitter.stop()
        self.assertTrue(self.splitter.is_stopped)
        self.assertFalse(self.splitter.is_paused)

    def test_reset(self):
        """Test reset functionality."""
        self.splitter.pause()
        self.splitter.stop()

        self.splitter.reset()
        self.assertFalse(self.splitter.is_paused)
        self.assertFalse(self.splitter.is_stopped)

    def test_split_single_file_basic(self):
        """Test basic single file splitting."""
        # Create a test document with delimiters
        paragraphs = [
            "Chapter 1: Introduction",
            "This is the first chapter.",
            "***",
            "Chapter 2: Methods",
            "This is the second chapter.",
            "***",
            "Chapter 3: Results",
            "This is the third chapter."
        ]

        input_file = self.create_test_document("test.docx", paragraphs)
        output_dir = os.path.join(self.test_dir, "output")

        # Split the document
        output_files = self.splitter.split_single_file(input_file, output_dir)

        # Verify results
        self.assertEqual(len(output_files), 3)
        self.assertTrue(all(os.path.exists(f) for f in output_files))

        # Verify first section content
        doc1 = Document(output_files[0])
        texts = [p.text for p in doc1.paragraphs]
        self.assertIn("Chapter 1: Introduction", texts)
        self.assertIn("This is the first chapter.", texts)
        self.assertNotIn("***", texts)

    def test_split_single_file_no_delimiters(self):
        """Test splitting file with no delimiters."""
        paragraphs = [
            "Chapter 1: Introduction",
            "This is the only chapter."
        ]

        input_file = self.create_test_document("test.docx", paragraphs)
        output_dir = os.path.join(self.test_dir, "output")

        # Split the document
        output_files = self.splitter.split_single_file(input_file, output_dir)

        # Should return empty list when no delimiters found
        self.assertEqual(len(output_files), 0)

    def test_split_single_file_invalid_input(self):
        """Test splitting with invalid input file."""
        with self.assertRaises(FileNotFoundError):
            self.splitter.split_single_file(
                "nonexistent.docx",
                self.test_dir
            )

    def test_split_single_file_invalid_extension(self):
        """Test splitting with invalid file extension."""
        # Create a text file
        txt_file = os.path.join(self.test_dir, "test.txt")
        with open(txt_file, 'w') as f:
            f.write("Not a Word document")

        with self.assertRaises(ValueError):
            self.splitter.split_single_file(txt_file, self.test_dir)

    def test_split_with_progress_callback(self):
        """Test splitting with progress callback."""
        paragraphs = [
            "Chapter 1",
            "***",
            "Chapter 2"
        ]

        input_file = self.create_test_document("test.docx", paragraphs)
        output_dir = os.path.join(self.test_dir, "output")

        progress_updates = []

        def progress_callback(percent, message):
            progress_updates.append((percent, message))

        self.splitter.split_single_file(
            input_file,
            output_dir,
            progress_callback=progress_callback
        )

        # Verify progress updates were made
        self.assertGreater(len(progress_updates), 0)
        self.assertEqual(progress_updates[-1][0], 100)

    def test_find_split_indices(self):
        """Test finding split indices in document."""
        paragraphs = [
            "Chapter 1",
            "***",
            "Chapter 2",
            "***",
            "Chapter 3"
        ]

        input_file = self.create_test_document("test.docx", paragraphs)
        doc = Document(input_file)

        indices = self.splitter._find_split_indices(doc)

        self.assertEqual(len(indices), 2)
        self.assertEqual(indices, [1, 3])

    def test_custom_delimiter(self):
        """Test splitting with custom delimiter."""
        self.splitter.delimiter = "###"

        paragraphs = [
            "Section 1",
            "###",
            "Section 2",
            "###",
            "Section 3"
        ]

        input_file = self.create_test_document("test.docx", paragraphs)
        output_dir = os.path.join(self.test_dir, "output")

        output_files = self.splitter.split_single_file(input_file, output_dir)

        self.assertEqual(len(output_files), 3)

    def test_batch_processing_empty_directory(self):
        """Test batch processing with empty directory."""
        empty_dir = os.path.join(self.test_dir, "empty")
        os.makedirs(empty_dir)

        output_dir = os.path.join(self.test_dir, "output")

        results = self.splitter.split_batch(empty_dir, output_dir)

        self.assertEqual(len(results['success']), 0)
        self.assertEqual(len(results['failed']), 0)

    def test_batch_processing_multiple_files(self):
        """Test batch processing with multiple files."""
        # Create input directory with multiple files
        input_dir = os.path.join(self.test_dir, "input")
        os.makedirs(input_dir)

        # Create multiple test documents
        for i in range(3):
            paragraphs = [
                f"File {i+1} - Chapter 1",
                "***",
                f"File {i+1} - Chapter 2"
            ]
            doc = Document()
            for text in paragraphs:
                doc.add_paragraph(text)
            doc.save(os.path.join(input_dir, f"file_{i+1}.docx"))

        output_dir = os.path.join(self.test_dir, "output")

        results = self.splitter.split_batch(input_dir, output_dir)

        self.assertEqual(len(results['success']), 3)
        self.assertEqual(len(results['failed']), 0)

    def test_batch_processing_with_failures(self):
        """Test batch processing with some failures."""
        input_dir = os.path.join(self.test_dir, "input")
        os.makedirs(input_dir)

        # Create one valid document
        paragraphs = ["Chapter 1", "***", "Chapter 2"]
        doc = Document()
        for text in paragraphs:
            doc.add_paragraph(text)
        doc.save(os.path.join(input_dir, "valid.docx"))

        # Create one invalid document (corrupted)
        invalid_file = os.path.join(input_dir, "invalid.docx")
        with open(invalid_file, 'w') as f:
            f.write("Not a valid docx file")

        output_dir = os.path.join(self.test_dir, "output")

        results = self.splitter.split_batch(input_dir, output_dir)

        self.assertEqual(len(results['success']), 1)
        self.assertEqual(len(results['failed']), 1)

    def test_output_directory_creation(self):
        """Test that output directory is created if it doesn't exist."""
        paragraphs = ["Chapter 1", "***", "Chapter 2"]
        input_file = self.create_test_document("test.docx", paragraphs)

        output_dir = os.path.join(self.test_dir, "new_output_dir")
        self.assertFalse(os.path.exists(output_dir))

        self.splitter.split_single_file(input_file, output_dir)

        self.assertTrue(os.path.exists(output_dir))

    def test_multiple_consecutive_delimiters(self):
        """Test handling of multiple consecutive delimiters."""
        paragraphs = [
            "Chapter 1",
            "***",
            "***",
            "Chapter 2"
        ]

        input_file = self.create_test_document("test.docx", paragraphs)
        output_dir = os.path.join(self.test_dir, "output")

        output_files = self.splitter.split_single_file(input_file, output_dir)

        # Should create sections even with consecutive delimiters
        self.assertGreater(len(output_files), 0)

    def test_delimiter_at_end(self):
        """Test handling delimiter at end of document."""
        paragraphs = [
            "Chapter 1",
            "***",
            "Chapter 2",
            "***"
        ]

        input_file = self.create_test_document("test.docx", paragraphs)
        output_dir = os.path.join(self.test_dir, "output")

        output_files = self.splitter.split_single_file(input_file, output_dir)

        # Should handle delimiter at end gracefully
        self.assertGreater(len(output_files), 0)


class TestWordSplitterIntegration(unittest.TestCase):
    """Integration tests for WordSplitter."""

    def setUp(self):
        """Set up test fixtures."""
        self.test_dir = tempfile.mkdtemp()
        self.splitter = WordSplitter()

    def tearDown(self):
        """Clean up test fixtures."""
        if os.path.exists(self.test_dir):
            shutil.rmtree(self.test_dir)

    def test_full_workflow(self):
        """Test complete workflow from input to output."""
        # Create a realistic test document
        doc = Document()
        doc.add_paragraph("Book Title", style='Title')
        doc.add_paragraph("Front Matter")
        doc.add_paragraph("This is the front matter.")
        doc.add_paragraph("***")
        doc.add_paragraph("Chapter 1: Introduction", style='Heading 1')
        doc.add_paragraph("This is the introduction.")
        doc.add_paragraph("***")
        doc.add_paragraph("Chapter 2: Methods", style='Heading 1')
        doc.add_paragraph("This is the methods section.")

        input_file = os.path.join(self.test_dir, "book.docx")
        doc.save(input_file)

        output_dir = os.path.join(self.test_dir, "chapters")

        # Split the document
        output_files = self.splitter.split_single_file(input_file, output_dir)

        # Verify
        self.assertEqual(len(output_files), 3)

        # Verify each file is valid and contains expected content
        for output_file in output_files:
            self.assertTrue(os.path.exists(output_file))
            # Should be able to open as Word document
            doc = Document(output_file)
            self.assertGreater(len(doc.paragraphs), 0)


if __name__ == '__main__':
    unittest.main()
